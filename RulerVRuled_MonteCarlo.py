"""
Ruler vs Ruled - Monte Carlo Simulation and Analysis Pipeline
Created by Larzsolice Aarend (find me on Medium)
Coded with assistance from Gemini

-------------------------------------------------------------

Required Third-Party Libraries:
- pandas (Data manipulation and Excel export)
- numpy (Numerical operations)
- scipy (Kernel Density Estimation for clustering)
- matplotlib & seaborn (Data visualization)
- scikit-learn (PCA and preprocessing)
- python-docx (Automated Word report generation)
- openpyxl (Excel file generation, utilized by pandas under the hood)

Simulation Overview:
- Simulation 1: Stochastic Baseline.
- Simulation 2: Blind Control mapped over fixed starting Quadrants (DD, DH, HD, HH, RR).
- Simulation 3 & 4: Neutral Players (Base) vs Neutral Players (Tit-For-Tat) across 5 Archetypes.
- Simulation 5 & 6: Aggressive Ruler (Base) vs Aggressive Ruler (Tit-For-Tat).
- Simulation 7 & 8: Aggressive Ruled (Base) vs Aggressive Ruled (Tit-For-Tat).
"""

if __name__ == "__main__": print("Initializing simulation environment...", flush=True)

# ==========================================
# SIMULATION PARAMETERS 
# ==========================================
PERCEIVED_VALUE = 10.0  # Value of a society
MAX_COST_PER_STEP = 1.0 # Per player
MAX_STEPS = 1000        # Per simulation
N_REPS = 1000           # Number of repetitions for each simulated situation
RECORD_HISTORY = False  # Slow and resource intensive, use on smaller N_REPS only


# ==========================================
# SLOW IMPORT LOADING
# ==========================================
n_import_chunks = 8

def print_import_progress(step, name):
    bar_length = 30
    filled = int(bar_length * step / n_import_chunks)
    bar = '#' * filled + '-' * (bar_length - filled)
    percent = (step / n_import_chunks) * 100
    if name == "":
        print(f"\rLoading core libraries: [{bar}] {percent:.0f}%                 \n", end="", flush=True)
    else:
        print(f"\rLoading core libraries: [{bar}] {percent:.0f}% ({name})     ", end="", flush=True)

# Fast imports
if __name__ == "__main__": print_import_progress(1, "Fast Imports")
import os
import shutil
import random
import zipfile
import multiprocessing
import time

# Slow imports
if __name__ == "__main__": print_import_progress(2, "pandas")
import pandas as pd

if __name__ == "__main__": print_import_progress(3, "numpy")
import numpy as np

if __name__ == "__main__": print_import_progress(4, "matplotlib")
import matplotlib.pyplot as plt
from matplotlib.patches import Patch

if __name__ == "__main__": print_import_progress(5, "seaborn")
import seaborn as sns

if __name__ == "__main__": print_import_progress(6, "scipy")
from scipy.stats import gaussian_kde

if __name__ == "__main__": print_import_progress(7, "sklearn")
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler

if __name__ == "__main__": print_import_progress(8, "python-docx")
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

if __name__ == "__main__": 
    print_import_progress(n_import_chunks, "")

# ==========================================
# 1. CONFIGURATION
# ==========================================

# KDE Settings
KDE_BW_ADJUST = "dynamic" 
KDE_TOP_N_PEAKS = 10 
KDE_OPTIMISATION_STEP = 0.01 
KDE_EXCLUDE_TOP_N_PEAKS = 5 
KDE_PROGRESSBAR_CHUNKSIZE = 300

# Graph Aesthetics
PLOT_DOT_SIZE = 10
PLOT_ALPHA = 0.5
CROSS_ALPHA = 1.0
CROSS_SIZE_RATIO = 0.90  
DIAMOND_ALPHA = 1.0
DIAMOND_SIZE_RATIO = 0.90
LINEWIDTH = 1
PLOT_PALETTE_PCA = "tab20"
PLOT_PALETTE_UNIFIED = "tab10"

# Performance Graph Settings
N_PERFORMERS = 5 # Number of top/bottom performing states to report and plot

# Target Filenames
FILE_SUMMARY = "results_summary.xlsx"
FILE_HISTORY = "results_history.xlsx"
FILE_PCA_FEATURES = "results_pca_features.xlsx"
FILE_PARAMS = "params.txt"
FILE_REPORT = "Report.docx"

OUTPUT_PREFIX = f"RulerVsRuled_V{PERCEIVED_VALUE}x{N_REPS}"
TEMP_DIR = "temp_export"
NUM_WORKERS = max(10, multiprocessing.cpu_count())  

# Custom Pair-Matched Palette for 4-Panel Performance Graphs (Control Last)
SIM_PALETTE = {
    'Sim 3 Neutral': '#1f77b4',        # Dark Blue
    'Sim 4 Neutral TFT': '#aec7e8',    # Light Blue
    'Sim 5 Agg Ruler': '#ff7f0e',      # Dark Orange
    'Sim 6 Agg Ruler TFT': '#ffbb78',  # Light Orange
    'Sim 7 Agg Ruled': '#d62728',      # Dark Red
    'Sim 8 Agg Ruled TFT': '#ff9896',  # Light Red
    'Sim 2 Control': '#7f7f7f'         # Grey (Placed last to ensure colors align with specifications)
}

_task_queue = None
_result_queue = None
_workers = []

# ==========================================
# 2. PERSISTENT MULTIPROCESSING ENGINE & CUSTOM CLASSES
# ==========================================
class CustomKDE(gaussian_kde):
    def __init__(self, dataset, bw_factor):
        self.bw_factor = bw_factor
        super().__init__(dataset)
    def covariance_factor(self):
        return super().covariance_factor() * self.bw_factor

def worker_loop(task_q, result_queue):
    import matplotlib
    matplotlib.use('Agg')
    while True:
        try:
            job = task_q.get()
            if job is None:  
                break
            job_id, command, payload = job
            
            if command == 'SIMULATE':
                sim_type, case, archetype = payload
                res, hist = run_game(sim_type, case, archetype)
                result_queue.put((job_id, (res, hist)))
            elif command == 'KDE_STEP':
                x, bw, grid_min, grid_max = payload
                try:
                    kde = CustomKDE(x, bw)
                    grid = np.linspace(grid_min, grid_max, 1000)
                    density = kde.evaluate(grid)
                    valleys_count = sum(1 for i in range(1, len(grid) - 1) if density[i] < density[i-1] and density[i] < density[i+1])
                    max_peak = np.max(density)
                    out = (bw, valleys_count, max_peak)
                except Exception:
                    out = (bw, 0, float('inf'))
                result_queue.put((job_id, out))
            elif command == 'PLOT_CASE':
                case, df_slice, features_for_pca, target_dir = payload
                out_path = render_single_case_pca(case, df_slice, features_for_pca, target_dir)
                result_queue.put((job_id, out_path))
            elif command == 'PLOT_CAT':
                cat, df_slice, features_for_pca, target_dir = payload
                out_path = render_single_category_pca(cat, df_slice, features_for_pca, target_dir)
                result_queue.put((job_id, out_path))
        except Exception as e:
            result_queue.put((job_id, e))

def init_persistent_workers():
    global _task_queue, _result_queue, _workers
    if not _workers:
        print(f"Initializing {NUM_WORKERS} persistent background worker processes...")
        _task_queue = multiprocessing.Queue()
        _result_queue = multiprocessing.Queue()
        for i in range(NUM_WORKERS):
            p = multiprocessing.Process(target=worker_loop, args=(_task_queue, _result_queue), name=f"PersistentWorker-{i+1}")
            p.daemon = True
            p.start()
            _workers.append(p)

def shutdown_persistent_workers():
    global _task_queue, _workers
    if _workers:
        print("Signaling background workers to shut down cleanly...")
        for _ in range(len(_workers)):
            _task_queue.put(None)
        for p in _workers:
            p.join(timeout=2.0)
        _workers = []

def run_parallel_jobs(jobs, desc="Progress"):
    global _task_queue, _result_queue
        
    # Initial progress display variables
    total_jobs = len(jobs)
    bar_length = 30
    bar_chars = '-' * bar_length
    filled_length = 0
    percentage = 0 * 100
    completed = 0
    
    for idx, (command, payload) in enumerate(jobs):
        _task_queue.put((idx, command, payload))
        
    results = [None] * total_jobs
    completed = 0
    start_time = time.time()

    print(f"\r{desc}: [{bar_chars}] 0/{total_jobs} (0.0%) | Elapsed: 00:00:00 | Speed: 0.0 jobs/s", end="", flush=True)
    
    while completed < total_jobs:
        try:
            job_id, out = _result_queue.get(timeout=1)
            results[job_id] = out
            completed += 1
            if completed % KDE_PROGRESSBAR_CHUNKSIZE == 0 or completed == total_jobs or total_jobs <= 100:
                filled_length = int(round(bar_length * completed / total_jobs))
                bar_chars = '#' * filled_length + '-' * (bar_length - filled_length)
                percentage = (completed / total_jobs) * 100
        except:
            pass
        
        if completed % KDE_PROGRESSBAR_CHUNKSIZE == 0 or completed == total_jobs or total_jobs <= 100:
            elapsed = time.time() - start_time
            elapsed_str = time.strftime("%H:%M:%S", time.gmtime(elapsed))
            speed = completed / elapsed if elapsed > 0 else 0

            print(f"\r{desc}: [{bar_chars}] {completed}/{total_jobs} ({percentage:.1f}%) | Elapsed: {elapsed_str} | Speed: {speed:.3f} jobs/s", end="", flush=True)
    
    print()
    return results

# ==========================================
# 3. CORE SIMULATION ENGINE
# ==========================================
def run_game(sim_type, case=None, archetype=None):
    random.seed()
    C_R, C_D = 0.0, 0.0
    history = {}
    cum_v_res = 0.0
    
    prob_r_list, prob_d_list = [], []
    cost_r_list, cost_d_list = [], []
    pen_r_list, pen_d_list = [], []
    
    if case == 'DD':
        case_p_R, case_p_D = random.uniform(0, 0.5), random.uniform(0, 0.5)
    elif case == 'DH':
        case_p_R, case_p_D = random.uniform(0, 0.5), random.uniform(0.5, 1.0)
    elif case == 'HD':
        case_p_R, case_p_D = random.uniform(0.5, 1.0), random.uniform(0, 0.5)
    elif case == 'HH':
        case_p_R, case_p_D = random.uniform(0.5, 1.0), random.uniform(0.5, 1.0)
    else: # 'RR' or 'None' (Random)
        case_p_R, case_p_D = random.uniform(0, 1.0), random.uniform(0, 1.0)

    p_R_base, p_D_base = case_p_R, case_p_D
    
    if sim_type in [3, 4]:
        p_R_base, p_D_base = random.uniform(0.0, 1.0), random.uniform(0.0, 1.0)
    elif sim_type in [5, 6]:
        p_R_base, p_D_base = random.uniform(0.50001, 1.0), random.uniform(0.0, 1.0)
    elif sim_type in [7, 8]:
        p_R_base, p_D_base = random.uniform(0.0, 1.0), random.uniform(0.50001, 1.0)

    p_R_hawk_if_H = random.uniform(0.50001, 1.0)  
    p_R_hawk_if_D = random.uniform(0.0, 0.49999)  
    p_D_hawk_if_H = random.uniform(0.50001, 1.0)  
    p_D_hawk_if_D = random.uniform(0.0, 0.49999)  

    step = 1
    dc_R = random.uniform(0, MAX_COST_PER_STEP)
    dc_D = random.uniform(0, MAX_COST_PER_STEP)
    
    contrib_R, contrib_D = dc_R, -dc_D
    
    cost_r_list.append(contrib_R / 2.0)
    cost_d_list.append(contrib_D / 2.0)
    prob_r_list.append(1.0)
    prob_d_list.append(0.0)
    pen_r_list.append(1.0)
    pen_d_list.append(1.0)
    
    prev_act_R, prev_act_D = 'H', 'D'
    
    delta_C = (contrib_R + contrib_D) / 2.0
    C_R, C_D = max(0.0, C_R + delta_C), max(0.0, C_D + delta_C)
    
    v_res = PERCEIVED_VALUE - (C_R + C_D)
    cum_v_res += max(0.0, v_res)
    
    if sim_type >= 3:
        next_penalty = max(0.0, v_res) / PERCEIVED_VALUE
        if archetype == 'Equals':
            next_pen_R, next_pen_D = next_penalty, next_penalty
        elif archetype == 'Cowards':
            next_pen_R, next_pen_D = next_penalty ** 2, next_penalty ** 2
        elif archetype == 'Fools':
            next_pen_R, next_pen_D = next_penalty ** 0.5, next_penalty ** 0.5
        elif archetype == 'Brinksmen':
            next_pen_R, next_pen_D = next_penalty, 1.0 - (1.0 - next_penalty) / 3.0
        elif archetype == 'Tyrants':
            next_pen_R, next_pen_D = 1.0 - (1.0 - next_penalty) / 3.0, next_penalty
        else:
            next_pen_R, next_pen_D = 1.0, 1.0
    else:
        next_pen_R, next_pen_D = 1.0, 1.0

    while (C_R + C_D) < PERCEIVED_VALUE and step < MAX_STEPS:
        step += 1
        v_res = PERCEIVED_VALUE - (C_R + C_D)
        
        if sim_type >= 3:
            penalty = v_res / PERCEIVED_VALUE
            if archetype == 'Equals':
                mul_R, mul_D = penalty, penalty
            elif archetype == 'Cowards':
                mul_R, mul_D = penalty ** 2, penalty ** 2
            elif archetype == 'Fools':
                mul_R, mul_D = penalty ** 0.5, penalty ** 0.5
            elif archetype == 'Brinksmen':
                mul_R, mul_D = penalty, 1.0 - (1.0 - penalty)/3.0
            elif archetype == 'Tyrants':
                mul_R, mul_D = 1.0 - (1.0 - penalty)/3.0, penalty
            else:
                mul_R, mul_D = 1.0, 1.0
        else:
            mul_R, mul_D = 1.0, 1.0

        if sim_type == 1:
            p_R_curr = random.uniform(0, 1.0)
            p_D_curr = random.uniform(0, 1.0)
        elif sim_type == 2:
            p_R_curr = p_R_base
            p_D_curr = p_D_base
        else:
            if step == 2:
                base_R, base_D = p_R_base, p_D_base
            else:
                if sim_type in [3, 5, 7]:
                    base_R, base_D = p_R_base, p_D_base
                elif sim_type == 4:
                    base_R = p_R_hawk_if_H if prev_act_D == 'H' else p_R_hawk_if_D
                    base_D = p_D_hawk_if_H if prev_act_R == 'H' else p_D_hawk_if_D
                elif sim_type == 6:
                    base_R = p_R_base
                    base_D = p_D_hawk_if_H if prev_act_R == 'H' else p_D_hawk_if_D
                elif sim_type == 8:
                    base_R = p_R_hawk_if_H if prev_act_D == 'H' else p_R_hawk_if_D
                    base_D = p_D_base
            
            p_R_curr = max(0.01, base_R * mul_R)
            p_D_curr = max(0.01, base_D * mul_D)

        prob_r_list.append(p_R_curr)
        prob_d_list.append(p_D_curr)
        pen_r_list.append(mul_R)
        pen_d_list.append(mul_D)

        act_R = 'H' if random.random() < p_R_curr else 'D'
        act_D = 'H' if random.random() < p_D_curr else 'D'
        
        prev_act_R, prev_act_D = act_R, act_D
        
        dc_R = random.uniform(0, MAX_COST_PER_STEP)
        dc_D = random.uniform(0, MAX_COST_PER_STEP)
        
        contrib_R = dc_R if act_R == 'H' else -dc_R
        contrib_D = dc_D if act_D == 'H' else -dc_D
        
        cost_r_list.append(contrib_R / 2.0)
        cost_d_list.append(contrib_D / 2.0)
        
        delta_C = (contrib_R + contrib_D) / 2.0
        
        C_R = max(0.0, C_R + delta_C)
        C_D = max(0.0, C_D + delta_C)
        
        v_res = max(0.0, PERCEIVED_VALUE - (C_R + C_D))
        cum_v_res += v_res
        
    if step == MAX_STEPS:
        term_status = "Survived"
    else:
        if_r_lost = contrib_R > contrib_D
        if if_r_lost:
            term_status = "Revolution"
        else:
            term_status = "Crackdown"
            
    result = {
        'Sim_Type': sim_type,
        'Case': case if case else 'Random',
        'Archetype': archetype if archetype else 'None',
        'Init_P_R': p_R_base,
        'Init_P_D': p_D_base,
        'Steps': step,
        'Cum_V_res': cum_v_res,
        'Terminal_Status': term_status,
        'Avg_P_R': np.mean(prob_r_list),
        'Avg_P_D': np.mean(prob_d_list),
        'Avg_C_R': np.mean(cost_r_list),
        'Avg_C_D': np.mean(cost_d_list),
        'Avg_Pen_R': np.mean(pen_r_list),
        'Avg_Pen_D': np.mean(pen_d_list)
    }
    return result, history

def _run_game_task_wrapper(args):
    sim_type, case, archetype = args
    return run_game(sim_type, case, archetype)

def run_all_simulations():
    cases = ['DD', 'DH', 'HD', 'HH', 'RR']
    archetypes = ['Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    jobs = []
    
    # Sim 1 (10x N_REPS)
    for _ in range(N_REPS * 10):
        jobs.append(('SIMULATE', (1, None, None)))
        
    # Sim 2 (N_REPS per case)
    for case in cases:
        for _ in range(N_REPS):
            jobs.append(('SIMULATE', (2, case, None)))
            
    # Sims 3 through 8 (Standardized identical simulation volume per strategic group)
    # Passed None for case here to strictly isolate 'RR' to Sim 2
    for sim_type in range(3, 9):
        for arch in archetypes:
            for _ in range(N_REPS * 5): 
                jobs.append(('SIMULATE', (sim_type, None, arch)))
                
    total_tasks = len(jobs)
    print(f"Total simulations to execute: {total_tasks}")
    task_results = run_parallel_jobs(jobs, desc="Executing Simulations")
    
    results = []
    histories = {}
    for idx, (res, hist) in enumerate(task_results):
        results.append(res)
        if RECORD_HISTORY:
            histories[idx] = hist
            
    df_res = pd.DataFrame(results)
    df_hist = pd.DataFrame() 
    return df_res, df_hist

# ==========================================
# 4. KDE VALLEY CLUSTERING IMPLEMENTATION
# ==========================================
def _eval_bandwidth_task(args):
    x, bw, grid_min, grid_max = args
    try:
        kde = CustomKDE(x, bw)
        grid = np.linspace(grid_min, grid_max, 1000)
        density = kde.evaluate(grid)
        valleys_count = sum(1 for i in range(1, len(grid) - 1) if density[i] < density[i-1] and density[i] < density[i+1])
        max_peak = np.max(density)
        return (bw, valleys_count, max_peak)
    except Exception:
        return (bw, 0, float('inf'))

def find_optimal_bandwidth(x, steps=None, top_n=KDE_TOP_N_PEAKS):
    if steps is None:
        steps = np.arange(KDE_OPTIMISATION_STEP, 1.0 + KDE_OPTIMISATION_STEP, KDE_OPTIMISATION_STEP)
    if len(x) < 2 or np.all(x == x[0]) or np.var(x) < 1e-9:
        return 0.3, CustomKDE(x, 0.3)
        
    grid_min, grid_max = x[0], x[-1]
    
    if multiprocessing.current_process().name == 'MainProcess' and _task_queue is not None:
        jobs = [('KDE_STEP', (x, bw, grid_min, grid_max)) for bw in steps]
        raw_candidates = run_parallel_jobs(jobs, desc="Finding Optimal KDE Bandwidths")
    else:
        raw_candidates = [_eval_bandwidth_task((x, bw, grid_min, grid_max)) for bw in steps]
            
    candidates = [item for item in raw_candidates if item is not None]
    if not candidates:
        return 0.3, CustomKDE(x, 0.3)
        
    candidates.sort(key=lambda item: item[1], reverse=True)
    
    if len(candidates) > KDE_EXCLUDE_TOP_N_PEAKS:
        shortlist = candidates[KDE_EXCLUDE_TOP_N_PEAKS : KDE_EXCLUDE_TOP_N_PEAKS + top_n]
    else:
        shortlist = candidates[:top_n]
        
    if not shortlist:
        return 0.3, CustomKDE(x, 0.3)
        
    best_candidate = min(shortlist, key=lambda item: item[2])
    best_bw = best_candidate[0]
    return best_bw, CustomKDE(x, best_bw)

def kde_valley_clustering(data, bw_adjust=KDE_BW_ADJUST, _global=False):
    global global_actual_bw
    x = np.sort(data)
    if len(x) < 2 or np.all(x == x[0]) or np.var(x) < 1e-9:
        return np.zeros(len(data), dtype=int)
        
    actual_bw, kde = (bw_adjust, CustomKDE(x, bw_adjust)) if bw_adjust != "dynamic" else find_optimal_bandwidth(x, top_n=KDE_TOP_N_PEAKS)
    if _global == True:
        global_actual_bw = actual_bw

    grid = np.linspace(x[0], x[-1], 1000)
    density = kde.evaluate(grid)
    
    valleys = [grid[i] for i in range(1, len(grid) - 1) if density[i] < density[i-1] and density[i] < density[i+1]]
            
    if len(valleys) > 19:
        valley_densities = kde.evaluate(valleys)
        sorted_indices = np.argsort(valley_densities)
        valleys = sorted([valleys[idx] for idx in sorted_indices[:19]])
        
    return np.digitize(data, valleys)

def get_ordered_local_clusters(data_values, bw_adjust=KDE_BW_ADJUST, _global=False):
    if len(data_values) == 0:
        return np.array([])
    if np.all(data_values == data_values[0]) or np.var(data_values) < 1e-9:
        return np.zeros(len(data_values), dtype=int)
        
    raw_labels = kde_valley_clustering(data_values, bw_adjust=bw_adjust, _global=_global)
    df_temp = pd.DataFrame({'val': data_values, 'raw_lbl': raw_labels})
    means = df_temp.groupby('raw_lbl')['val'].mean().sort_values(ascending=False)
    mapping = {old_id: new_id + 1 for new_id, old_id in enumerate(means.index)}
    return df_temp['raw_lbl'].map(mapping).values

# ==========================================
# 5. DATA POST-PROCESSING & MULTIDIMENSIONAL PCA
# ==========================================
def process_data(df):
    print("Performing Global KDE Valley Clustering...")
    df['Cluster'] = [f"{int(c):02d}" for c in get_ordered_local_clusters(df['Cum_V_res'].values, bw_adjust=KDE_BW_ADJUST, _global=True)]
    
    sorted_unique_clusters = sorted(df['Cluster'].unique())
    top_5_clusters = sorted_unique_clusters[:5] if len(sorted_unique_clusters) >= 5 else sorted_unique_clusters
    
    survival_rates = df.groupby('Cluster')['Steps'].min()
    full_survival_clusters = survival_rates[survival_rates == MAX_STEPS].index
    bottom_5_survival = sorted(full_survival_clusters)[:5] if len(full_survival_clusters) > 0 else []
    termination_clusters = survival_rates[survival_rates < MAX_STEPS].index
    
    print("Performing Global Multi-Dimensional PCA...")
    features_for_pca = ['Avg_P_R', 'Avg_P_D', 'Avg_C_R', 'Avg_C_D', 'Avg_Pen_R', 'Avg_Pen_D']
    x_matrix = df[features_for_pca].values
    x_scaled = StandardScaler().fit_transform(x_matrix)
    
    pca = PCA(n_components=2)
    principal_components = pca.fit_transform(x_scaled)
    df['PCA_1'] = principal_components[:, 0]
    df['PCA_2'] = principal_components[:, 1]
    
    all_export_features = ['Avg_P_R', 'Avg_P_D', 'Avg_C_R', 'Avg_C_D', 'Avg_Pen_R', 'Avg_Pen_D', 'Cum_V_res']
    df_pca_features = df[['Sim_Type', 'Case', 'Archetype', 'Cluster'] + all_export_features + ['PCA_1', 'PCA_2']].copy()
    
    return df, df_pca_features, top_5_clusters, bottom_5_survival, termination_clusters


# ==========================================
# CORE PLOT RENDERING FUNCTIONS
# ==========================================
def render_single_case_pca(case, df_slice, features_for_pca, target_dir):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
    sns.set_theme(style="whitegrid")
    
    plt.figure(figsize=(8, 6))
    if not df_slice.empty and len(df_slice) >= 5:
        x_local = df_slice[features_for_pca].values
        stds = np.std(x_local, axis=0)
        stds[stds == 0] = 1.0
        x_scaled = (x_local - np.mean(x_local, axis=0)) / stds
        
        pca_local = PCA(n_components=2)
        pcs_local = pca_local.fit_transform(x_scaled)
        
        df_slice = df_slice.copy()
        df_slice['Local_PCA_1'] = pcs_local[:, 0]
        df_slice['Local_PCA_2'] = pcs_local[:, 1]
        
        local_lbls = get_ordered_local_clusters(df_slice['Cum_V_res'].values, bw_adjust=KDE_BW_ADJUST)
        df_slice['Local_Cluster'] = [f"{int(l):02d}" for l in local_lbls]
        
        df_slice_sorted = df_slice.sort_values('Local_Cluster')
        unique_clusters = sorted(df_slice_sorted['Local_Cluster'].unique())
        
        # Plot circles with NO borders
        sns.scatterplot(
            x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
            data=df_slice_sorted, palette='tab20', alpha=PLOT_ALPHA, s=PLOT_DOT_SIZE, marker='o',
            edgecolor=None, linewidth=0
        )
        
        terminated_runs = df_slice_sorted[df_slice_sorted['Terminal_Status'] != 'Survived']
        if not terminated_runs.empty:
            sns.scatterplot(
                x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
                data=terminated_runs, palette='tab20', alpha=CROSS_ALPHA, s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, marker='x',
                legend=False, linewidth=LINEWIDTH
            )
            
        sorted_by_vres = df_slice_sorted.sort_values('Cum_V_res', ascending=False)
        k_1pct = max(1, int(len(sorted_by_vres) * 0.01))
        num_highlights = min(k_1pct, 100)
        
        # Highlight top 1% highest cumulative V_res scores with diamonds
        top_1pct_runs = sorted_by_vres.head(num_highlights)
        if not top_1pct_runs.empty:
            sns.scatterplot(
                x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
                data=top_1pct_runs, palette='tab20', alpha=DIAMOND_ALPHA, s=PLOT_DOT_SIZE * DIAMOND_SIZE_RATIO, marker='d',
                legend=False, edgecolor='black', linewidth=LINEWIDTH
            )
            
        # Highlight bottom 1% lowest cumulative V_res scores (Black crosses if terminated, transparent circles if survived)
        bottom_1pct_runs = sorted_by_vres.tail(num_highlights)
        bottom_terminated = bottom_1pct_runs[bottom_1pct_runs['Terminal_Status'] != 'Survived']
        bottom_survived = bottom_1pct_runs[bottom_1pct_runs['Terminal_Status'] == 'Survived']
        
        if not bottom_terminated.empty:
            plt.scatter(
                bottom_terminated['Local_PCA_1'], bottom_terminated['Local_PCA_2'],
                marker='x', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, color='black', alpha=CROSS_ALPHA,
                linewidths=LINEWIDTH, zorder=5
            )
            
        if not bottom_survived.empty:
            plt.scatter(
                bottom_survived['Local_PCA_1'], bottom_survived['Local_PCA_2'],
                marker='o', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, facecolors='none', edgecolors='black', 
                alpha=CROSS_ALPHA, linewidths=LINEWIDTH, zorder=5
            )
        
    plt.title(f"Isolated PCA & Local KDE: Case {case}")
    plt.xlabel("Principal Component 1 (Local Component)")
    plt.ylabel("Principal Component 2 (Local Component)")
    plt.legend(title='Local KDE Cluster', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.tight_layout()
    path = os.path.join(target_dir, f"pca_case_{case}.png")
    plt.savefig(path)
    plt.close()
    return path


def render_single_category_pca(cat, df_slice, features_for_pca, target_dir):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
    sns.set_theme(style="whitegrid")
    
    plt.figure(figsize=(8, 6))
    
    if cat == 'Sim_2_Control':
        title_label = "Sim 2 (Control) Baseline"
    else:
        parts = cat.split('_')
        s_num = parts[1]
        a_name = parts[2]
        sim_names = {
            '3': 'Sim 3 (Neutral)', '4': 'Sim 4 (Neutral TFT)',
            '5': 'Sim 5 (Agg Ruler)', '6': 'Sim 6 (Agg Ruler TFT)',
            '7': 'Sim 7 (Agg Ruled)', '8': 'Sim 8 (Agg Ruled TFT)'
        }
        title_label = f"{sim_names[s_num]} - {a_name}"
    
    if not df_slice.empty and len(df_slice) >= 5:
        x_local = df_slice[features_for_pca].values
        stds = np.std(x_local, axis=0)
        stds[stds == 0] = 1.0
        x_scaled = (x_local - np.mean(x_local, axis=0)) / stds
        
        pca_local = PCA(n_components=2)
        pcs_local = pca_local.fit_transform(x_scaled)
        
        df_slice = df_slice.copy()
        df_slice['Local_PCA_1'] = pcs_local[:, 0]
        df_slice['Local_PCA_2'] = pcs_local[:, 1]
        
        local_lbls = get_ordered_local_clusters(df_slice['Cum_V_res'].values, bw_adjust=KDE_BW_ADJUST)
        df_slice['Local_Cluster'] = [f"{int(l):02d}" for l in local_lbls]
        
        df_slice_sorted = df_slice.sort_values('Local_Cluster')
        unique_clusters = sorted(df_slice_sorted['Local_Cluster'].unique())
        
        # Plot circles with NO borders
        sns.scatterplot(
            x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
            data=df_slice_sorted, palette='tab20', alpha=PLOT_ALPHA, s=PLOT_DOT_SIZE, marker='o',
            edgecolor=None, linewidth=0
        )
        
        terminated_runs = df_slice_sorted[df_slice_sorted['Terminal_Status'] != 'Survived']
        if not terminated_runs.empty:
            sns.scatterplot(
                x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
                data=terminated_runs, palette='tab20', alpha=CROSS_ALPHA, s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, marker='x',
                legend=False, linewidth=LINEWIDTH
            )
            
        sorted_by_vres = df_slice_sorted.sort_values('Cum_V_res', ascending=False)
        k_1pct = max(1, int(len(sorted_by_vres) * 0.01))
        num_highlights = min(k_1pct, 100)
        
        # Highlight top 1% highest cumulative V_res scores with diamonds
        top_1pct_runs = sorted_by_vres.head(num_highlights)
        if not top_1pct_runs.empty:
            sns.scatterplot(
                x='Local_PCA_1', y='Local_PCA_2', hue='Local_Cluster', hue_order=unique_clusters,
                data=top_1pct_runs, palette='tab20', alpha=DIAMOND_ALPHA, s=PLOT_DOT_SIZE * DIAMOND_SIZE_RATIO, marker='d',
                legend=False, edgecolor='black', linewidth=LINEWIDTH
            )
            
        # Highlight bottom 1% lowest cumulative V_res scores
        bottom_1pct_runs = sorted_by_vres.tail(num_highlights)
        bottom_terminated = bottom_1pct_runs[bottom_1pct_runs['Terminal_Status'] != 'Survived']
        bottom_survived = bottom_1pct_runs[bottom_1pct_runs['Terminal_Status'] == 'Survived']
        
        if not bottom_terminated.empty:
            plt.scatter(
                bottom_terminated['Local_PCA_1'], bottom_terminated['Local_PCA_2'],
                marker='x', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, color='black', alpha=CROSS_ALPHA,
                linewidths=LINEWIDTH, zorder=5
            )
            
        if not bottom_survived.empty:
            plt.scatter(
                bottom_survived['Local_PCA_1'], bottom_survived['Local_PCA_2'],
                marker='o', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, facecolors='none', edgecolors='black', 
                alpha=CROSS_ALPHA, linewidths=LINEWIDTH, zorder=5
            )
        
    plt.title(f"Isolated PCA: {title_label}")
    plt.xlabel("Principal Component 1 (Local Component)")
    plt.ylabel("Principal Component 2 (Local Component)")
    plt.legend(title='Local KDE Cluster', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.tight_layout()
    path = os.path.join(target_dir, f"pca_cat_{cat}.png")
    plt.savefig(path)
    plt.close()
    return path


# ==========================================
# 6. VISUALIZATIONS
# ==========================================
def generate_graphs(df, top_5, bottom_5_surv, term_clusters, target_dir):
    os.makedirs(target_dir, exist_ok=True)
    images = []
    sns.set_theme(style="whitegrid")
    
    features_for_pca = ['Avg_P_R', 'Avg_P_D', 'Avg_C_R', 'Avg_C_D', 'Avg_Pen_R', 'Avg_Pen_D']
    
    print("Generating Unified Archetypes PCA Plot...")
    plt.figure(figsize=(10, 7))
    unique_archetypes = sorted(df[df['Archetype'] != 'None']['Archetype'].unique())
    sns.scatterplot(
        x='PCA_1', y='PCA_2', hue='Archetype', hue_order=unique_archetypes, style='Archetype',
        data=df[df['Archetype'] != 'None'], palette=PLOT_PALETTE_UNIFIED, alpha=PLOT_ALPHA, s=PLOT_DOT_SIZE, marker='o',
        edgecolor=None, linewidth=0
    )
    
    unified_terminated = df[(df['Archetype'] != 'None') & (df['Terminal_Status'] != 'Survived')]
    if not unified_terminated.empty:
        sns.scatterplot(
            x='PCA_1', y='PCA_2', hue='Archetype', hue_order=unique_archetypes,
            data=unified_terminated, palette=PLOT_PALETTE_UNIFIED, alpha=CROSS_ALPHA, s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, marker='x',
            legend=False, linewidth=LINEWIDTH
        )
        
    sorted_by_vres_global = df[df['Archetype'] != 'None'].sort_values('Cum_V_res', ascending=False)
    k_global = max(1, int(len(sorted_by_vres_global) * 0.01))
    num_global_highlights = min(k_global, 100)
    
    # Highlighting elite top performers
    top_1pct_global = sorted_by_vres_global.head(num_global_highlights)
    if not top_1pct_global.empty:
        sns.scatterplot(
            x='PCA_1', y='PCA_2', hue='Archetype', hue_order=unique_archetypes,
            data=top_1pct_global, palette=PLOT_PALETTE_UNIFIED, alpha=DIAMOND_ALPHA, s=PLOT_DOT_SIZE * DIAMOND_SIZE_RATIO, marker='d',
            legend=False, edgecolor='black', linewidth=LINEWIDTH
        )
        
    # Highlighting worst performers globally
    bottom_1pct_global = sorted_by_vres_global.tail(num_global_highlights)
    bottom_term_global = bottom_1pct_global[bottom_1pct_global['Terminal_Status'] != 'Survived']
    bottom_surv_global = bottom_1pct_global[bottom_1pct_global['Terminal_Status'] == 'Survived']
    
    if not bottom_term_global.empty:
        plt.scatter(
            bottom_term_global['PCA_1'], bottom_term_global['PCA_2'],
            marker='x', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, color='black', alpha=CROSS_ALPHA,
            linewidths=LINEWIDTH, zorder=5
        )
        
    if not bottom_surv_global.empty:
        plt.scatter(
            bottom_surv_global['PCA_1'], bottom_surv_global['PCA_2'],
            marker='o', s=PLOT_DOT_SIZE * CROSS_SIZE_RATIO, facecolors='none', edgecolors='black', 
            alpha=CROSS_ALPHA, linewidths=LINEWIDTH, zorder=5
        )
        
    plt.title("PCA of Simulations: National Archetypes Projected")
    plt.xlabel("Principal Component 1")
    plt.ylabel("Principal Component 2")
    plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.tight_layout()
    pca_uni_path = os.path.join(target_dir, "pca_unified_archetypes.png")
    plt.savefig(pca_uni_path)
    images.append(pca_uni_path)
    plt.close()

    # Parallel Rendering
    plot_jobs = []
    cases = ['DD', 'DH', 'HD', 'HH', 'RR']
    for case in cases:
        case_slice = df[df['Case'] == case].copy()
        plot_jobs.append(('PLOT_CASE', (case, case_slice, features_for_pca, target_dir)))
        
    archetypes = ['Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    
    cat_slice_ctrl = df[df['Sim_Type'] == 2].copy()
    plot_jobs.append(('PLOT_CAT', ('Sim_2_Control', cat_slice_ctrl, features_for_pca, target_dir)))
    
    for sim_type in range(3, 9):
        for arch in archetypes:
            cat_label = f"Sim_{sim_type}_{arch}"
            cat_slice = df[(df['Sim_Type'] == sim_type) & (df['Archetype'] == arch)].copy()
            plot_jobs.append(('PLOT_CAT', (cat_label, cat_slice, features_for_pca, target_dir)))
        
    rendered_paths_list = run_parallel_jobs(plot_jobs, desc="Rendering PCA Plots")
    
    rendered_paths = {}
    job_idx = 0
    for case in cases:
        rendered_paths[('case', case)] = rendered_paths_list[job_idx]
        job_idx += 1
        
    rendered_paths[('category', 'Sim_2_Control')] = rendered_paths_list[job_idx]
    job_idx += 1
    
    for sim_type in range(3, 9):
        for arch in archetypes:
            rendered_paths[('category', f"Sim_{sim_type}_{arch}")] = rendered_paths_list[job_idx]
            job_idx += 1

    for case in cases:
        images.append(rendered_paths[('case', case)])

    print("Generating Complete Cluster Abundance Plots for all simulations...")
    for sim_type in range(1, 9):
        df_sim = df[df['Sim_Type'] == sim_type].copy()
        local_lbls = get_ordered_local_clusters(df_sim['Cum_V_res'].values, bw_adjust=KDE_BW_ADJUST)
        df_sim['Local_Cluster'] = [f"{int(l):02d}" for l in local_lbls]
        df_sim_sorted = df_sim.sort_values('Local_Cluster')
        
        plt.figure(figsize=(12, 6))
        if sim_type == 1:
            try:
                ax = sns.countplot(x='Local_Cluster', data=df_sim_sorted, palette='viridis', hue='Local_Cluster')
                if ax.get_legend() is not None:
                    ax.get_legend().remove()
            except Exception:
                ax = sns.countplot(x='Local_Cluster', data=df_sim_sorted, palette='viridis')
            plt.title(f"Simulation {sim_type} Complete Cluster Abundance")
        elif sim_type == 2:
            sns.countplot(x='Local_Cluster', data=df_sim_sorted, color='skyblue')
            plt.title(f"Simulation {sim_type} Complete Cluster Abundance")
        else:
            sns.countplot(x='Local_Cluster', hue='Archetype', data=df_sim_sorted, palette='Set2')
            plt.title(f"Simulation {sim_type} Cluster Abundance by Archetype")
            plt.legend(title='Archetype', bbox_to_anchor=(1.05, 1), loc='upper left')
            
        plt.xlabel("Local KDE Cluster ID")
        plt.ylabel("Number of Simulations")
        plt.tight_layout()
        sim_path = os.path.join(target_dir, f"abundance_sim{sim_type}.png")
        plt.savefig(sim_path)
        images.append(sim_path)
        plt.close()

    def plot_stacked_terminations_no_survival(df_subset, group_col, title, filename, expected_categories=None, figsize=(10,6), rotation=0):
        if expected_categories is None:
            expected_categories = sorted(df_subset[group_col].dropna().unique())
            
        counts = pd.crosstab(df_subset[group_col], df_subset['Terminal_Status'])
        counts = counts.reindex(expected_categories, fill_value=0)
        for col in ['Survived', 'Revolution', 'Crackdown']:
            if col not in counts.columns: counts[col] = 0
                
        row_sums = counts.sum(axis=1) 
        props = counts.div(row_sums, axis=0).fillna(0) * 100
        props_to_plot = props[['Revolution', 'Crackdown']]
        counts_to_plot = counts[['Revolution', 'Crackdown']]
        
        plt.figure(figsize=figsize)
        ax = plt.gca()
        props_to_plot.plot(kind='bar', stacked=True, color=['#e74c3c', '#2980b9'], edgecolor='black', ax=ax)
        
        new_labels = []
        for cat in expected_categories:
            sz = row_sums[cat]
            new_labels.append(f"{cat}\n(N={sz})")
        ax.set_xticklabels(new_labels, rotation=rotation, ha='center' if rotation==0 else 'right')
        
        n_info = f"N = {row_sums.iloc[0]} per category" if len(row_sums.unique()) == 1 else "N varies (see labels)"
        plt.title(f"{title}\n({n_info} | Heights represent Total Failure %)")
        plt.xlabel(group_col)
        plt.ylabel("Proportion of Total State Failures (%)")
        
        legend_elements = [
            Patch(facecolor='#e74c3c', edgecolor='black', label='Revolution'),
            Patch(facecolor='#2980b9', edgecolor='black', label='Crackdown')
        ]
        ax.legend(handles=legend_elements, title="Termination Driver", bbox_to_anchor=(1.05, 1), loc='upper left')
        
        for c_idx, col in enumerate(props_to_plot.columns):
            for r_idx, row in enumerate(props_to_plot.index):
                val = counts_to_plot.loc[row, col]
                if val > 0:
                    y_pos = props_to_plot.iloc[r_idx, :c_idx].sum() + props_to_plot.loc[row, col] / 2.0
                    ax.text(r_idx, y_pos, f"{int(val)}", ha='center', va='center', color='white', fontweight='bold', fontsize=9)
        
        plt.ylim(0, max(props_to_plot.sum(axis=1).max() * 1.15, 15))
        plt.tight_layout()
        path = os.path.join(target_dir, filename)
        plt.savefig(path)
        plt.close()
        return path

    print("Generating Failure proportions by simulation type...")
    df_sim_labeled = df.copy()
    df_sim_labeled['Simulation'] = df_sim_labeled['Sim_Type'].apply(lambda x: f"Sim {x}")
    path_term_sim = plot_stacked_terminations_no_survival(
        df_sim_labeled, 'Simulation', "State Failures by Simulation Framework", 
        "termination_proportions_sim.png", expected_categories=[f"Sim {i}" for i in range(1, 9)]
    )
    images.append(path_term_sim)

    print("Generating Failure proportions by starting quadrant...")
    df_cases_only = df[(df['Sim_Type'] == 2) & df['Case'].isin(cases)].copy()
    path_term_case = plot_stacked_terminations_no_survival(
        df_cases_only, 'Case', "State Failures by Case Quadrant", 
        "termination_proportions_case.png", expected_categories=cases
    )
    images.append(path_term_case)

    print("Generating Failure proportions by strategic category (Multi-panel)...")
    df_cat = df[df['Sim_Type'] >= 2].copy()
    fig, axes = plt.subplots(2, 3, figsize=(20, 12), sharey=True)
    archetypes_for_plot = ['Control', 'Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    
    sim_labels = {
        2: 'Sim 2\nControl',
        3: 'Sim 3\nNeutral',
        4: 'Sim 4\nNeutral TFT',
        5: 'Sim 5\nAgg Ruler',
        6: 'Sim 6\nAgg Ruler TFT',
        7: 'Sim 7\nAgg Ruled',
        8: 'Sim 8\nAgg Ruled TFT'
    }
    
    for idx, (ax, arch) in enumerate(zip(axes.flatten(), archetypes_for_plot)):
        if arch == 'Control':
            df_sub = df_cat[df_cat['Sim_Type'] == 2].copy()
            expected_cats = [2]
        else:
            df_sub = df_cat[(df_cat['Sim_Type'] >= 3) & (df_cat['Archetype'] == arch)].copy()
            expected_cats = [3, 4, 5, 6, 7, 8]
            
        counts = pd.crosstab(df_sub['Sim_Type'], df_sub['Terminal_Status'])
        counts = counts.reindex(expected_cats, fill_value=0)
        
        for col in ['Survived', 'Revolution', 'Crackdown']:
            if col not in counts.columns: counts[col] = 0
            
        row_sums = counts.sum(axis=1)
        props = counts.div(row_sums, axis=0).fillna(0) * 100
        props_to_plot = props[['Revolution', 'Crackdown']]
        counts_to_plot = counts[['Revolution', 'Crackdown']]
        
        props_to_plot.plot(kind='bar', stacked=True, color=['#e74c3c', '#2980b9'], edgecolor='black', ax=ax, legend=False)
        
        new_labels = [f"{sim_labels[cat]}\n(N={row_sums[cat]})" for cat in expected_cats]
        ax.set_xticklabels(new_labels, rotation=0, ha='center')
        
        ax.set_title(f"Archetype: {arch}", fontweight='bold')
        ax.set_xlabel("")
        if ax in axes[:, 0]:
            ax.set_ylabel("Proportion of Total State Failures (%)")
            
        for c_idx, col in enumerate(props_to_plot.columns):
            for r_idx, row in enumerate(props_to_plot.index):
                val = counts_to_plot.loc[row, col]
                if val > 0:
                    y_pos = props_to_plot.iloc[r_idx, :c_idx].sum() + props_to_plot.loc[row, col] / 2.0
                    ax.text(r_idx, y_pos, f"{int(val)}", ha='center', va='center', color='white', fontweight='bold', fontsize=9)
                    
    legend_elements = [
        Patch(facecolor='#e74c3c', edgecolor='black', label='Revolution'),
        Patch(facecolor='#2980b9', edgecolor='black', label='Crackdown')
    ]
    fig.legend(handles=legend_elements, title="Termination Driver", bbox_to_anchor=(0.5, 1.05), loc='center', ncol=2, fontsize=12, title_fontsize=14)
    
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    path_term_cat = os.path.join(target_dir, "termination_proportions_category.png")
    plt.savefig(path_term_cat)
    plt.close()
    images.append(path_term_cat)

    return images, rendered_paths

def safe_barplot(ax, x, y, hue, data, palette, hue_order=None):
    try:
        sns.barplot(x=x, y=y, hue=hue, data=data, palette=palette, hue_order=hue_order, capsize=0.1, errorbar='sd', ax=ax)
    except Exception:
        try:
            sns.barplot(x=x, y=y, hue=hue, data=data, palette=palette, hue_order=hue_order, capsize=0.1, ci='sd', ax=ax)
        except Exception:
            sns.barplot(x=x, y=y, hue=hue, data=data, palette=palette, hue_order=hue_order, ax=ax)

# ==========================================
# 7. PERFORMANCE BAR GRAPHS (TOP/BOTTOM 5)
# ==========================================
def generate_performance_graphs(df, target_dir):
    os.makedirs(target_dir, exist_ok=True)
    perf_images = []
    
    archetypes = ['Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    cases = ['DD', 'DH', 'HD', 'HH', 'RR']
    
    custom_color_map = {
        'Sim 3 Neutral': '#1f77b4',        
        'Sim 4 Neutral TFT': '#aec7e8',    
        'Sim 5 Agg Ruler': '#ff7f0e',      
        'Sim 6 Agg Ruler TFT': '#ffbb78',  
        'Sim 7 Agg Ruled': '#d62728',      
        'Sim 8 Agg Ruled TFT': '#ff9896',  
        'Sim 2 Control': '#7f7f7f'         
    }
    
    for extremity in ['Top', 'Bottom']:
        data_records = []
        ascending_order = True if extremity == 'Bottom' else False
        
        for sim_type in range(3, 9):
            sim_label = {3:'Sim 3 Neutral', 4:'Sim 4 Neutral TFT', 5:'Sim 5 Agg Ruler', 6:'Sim 6 Agg Ruler TFT', 7:'Sim 7 Agg Ruled', 8:'Sim 8 Agg Ruled TFT'}[sim_type]
            for arch in archetypes:
                df_sub = df[(df['Sim_Type'] == sim_type) & (df['Archetype'] == arch)]
                top_perf = df_sub.sort_values('Cum_V_res', ascending=ascending_order).head(N_PERFORMERS)
                for _, row in top_perf.iterrows():
                    data_records.append({
                        'Archetype': arch, 'Simulation': sim_label,
                        'Ruler P_H': row['Avg_P_R'], 'Ruled P_H': row['Avg_P_D'],
                        'Ruler Cost': row['Avg_C_R'], 'Ruled Cost': row['Avg_C_D'],
                        'Cum_V_res': row['Cum_V_res']
                    })
                    
        df_ctrl = df[df['Sim_Type'] == 2]
        top_ctrl = df_ctrl.sort_values('Cum_V_res', ascending=ascending_order).head(N_PERFORMERS)
        for _, row in top_ctrl.iterrows():
            data_records.append({
                'Archetype': 'Control', 'Simulation': 'Sim 2 Control',
                'Ruler P_H': row['Avg_P_R'], 'Ruled P_H': row['Avg_P_D'],
                'Ruler Cost': row['Avg_C_R'], 'Ruled Cost': row['Avg_C_D'],
                'Cum_V_res': row['Cum_V_res']
            })
                    
        df_perf = pd.DataFrame(data_records)
        
        if not df_perf.empty:
            df_melt = pd.melt(
                df_perf, id_vars=['Archetype', 'Simulation'], 
                value_vars=['Ruler P_H', 'Ruled P_H', 'Ruler Cost', 'Ruled Cost'],
                var_name='Metric', value_name='Value'
            )
            
            fig, axes = plt.subplots(2, 2, figsize=(16, 12))
            metrics = ['Ruler P_H', 'Ruled P_H', 'Ruler Cost', 'Ruled Cost']
            x_order = archetypes + ['Control']
            hue_order = list(custom_color_map.keys())
            
            for ax, metric in zip(axes.flatten(), metrics):
                safe_barplot(ax, 'Archetype', 'Value', 'Simulation', df_melt[df_melt['Metric'] == metric], custom_color_map, hue_order=hue_order)
                ax.set_title(f"Average {metric}", fontweight='bold')
                ax.set_ylabel("Value (Mean +/- SD)")
                ax.set_xlabel("")
                if ax != axes[0, 1]:
                    if ax.get_legend() is not None:
                        ax.get_legend().remove()
                else:
                    ax.legend(title="Framework", bbox_to_anchor=(1.05, 1), loc='upper left')
                    
            plt.suptitle(f"Performance Metrics: 4-Panel Analysis ({extremity} {N_PERFORMERS} Performers)", fontsize=16, fontweight='bold')
            plt.tight_layout(rect=[0, 0, 0.9, 1])
            
            img_path = os.path.join(target_dir, f"perf_archetypes_4panel_{extremity.lower()}.png")
            plt.savefig(img_path)
            plt.close()
            perf_images.append(img_path)

    fig, axes = plt.subplots(1, 2, figsize=(20, 8), sharey=True)
    
    for idx, extremity in enumerate(['Top', 'Bottom']):
        ax = axes[idx]
        ascending_order = True if extremity == 'Bottom' else False
        pros_records = []
        
        for sim_type in range(3, 9):
            sim_label = {3:'Sim 3 Neutral', 4:'Sim 4 Neutral TFT', 5:'Sim 5 Agg Ruler', 6:'Sim 6 Agg Ruler TFT', 7:'Sim 7 Agg Ruled', 8:'Sim 8 Agg Ruled TFT'}[sim_type]
            for arch in archetypes:
                df_sub = df[(df['Sim_Type'] == sim_type) & (df['Archetype'] == arch)]
                top_perf = df_sub.sort_values('Cum_V_res', ascending=ascending_order).head(N_PERFORMERS)
                for _, row in top_perf.iterrows():
                    pros_records.append({'Archetype': arch, 'Simulation': sim_label, 'Cum_V_res': row['Cum_V_res']})
                    
        df_ctrl = df[df['Sim_Type'] == 2]
        top_ctrl = df_ctrl.sort_values('Cum_V_res', ascending=ascending_order).head(N_PERFORMERS)
        for _, row in top_ctrl.iterrows():
            pros_records.append({'Archetype': 'Control', 'Simulation': 'Sim 2 Control', 'Cum_V_res': row['Cum_V_res']})
            
        df_pros = pd.DataFrame(pros_records)
        if not df_pros.empty:
            safe_barplot(ax, 'Archetype', 'Cum_V_res', 'Simulation', df_pros, custom_color_map, hue_order=list(custom_color_map.keys()))
            ax.set_title(f"{extremity} {N_PERFORMERS} Performers Prosperity", fontweight='bold')
            ax.set_ylabel("Cumulative Perceived Value (Cum_V_res)" if idx == 0 else "")
            ax.set_xlabel("")
            if idx == 0:
                if ax.get_legend() is not None:
                    ax.get_legend().remove()
            else:
                ax.legend(title="Framework", bbox_to_anchor=(1.05, 1), loc='upper left')

    plt.suptitle("Prosperity Scores (Cumulative Vresidual) Comparison Across Simulations", fontsize=16, fontweight='bold')
    plt.tight_layout(rect=[0, 0, 0.9, 1])
    pros_path = os.path.join(target_dir, "perf_consolidated_prosperity.png")
    plt.savefig(pros_path)
    plt.close()
    perf_images.append(pros_path)
    
    return perf_images

# ==========================================
# 8. REPORT & ARCHIVE GENERATION
# ==========================================
def generate_report(images, rendered_paths, perf_images, target_dir, df):
    doc = docx.Document()
    
    # Modify global styles
    if 'Normal' in doc.styles:
        style_normal = doc.styles['Normal']
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_normal.font.name = 'Candara'
        
    for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    def add_centered_image(doc_obj, img_path, width):
        p = doc_obj.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(img_path, width=width)

    doc.add_heading('Ruler vs Ruled: Monte Carlo Simulation Report', 0)
    doc.add_paragraph("Ruler vs Ruled model created by Larzsolice Aarend (find me on Medium).")
    
    doc.add_heading('1. Introduction & Game Overview', level=1)
    doc.add_paragraph("This report presents the findings from an advanced Monte Carlo simulation of the 'Ruler vs. Ruled' dynamic. The model is structured as an escalating, multi-turn game of Chicken, where two actors (the Ruler and the Ruled) interact over a maximum of 1,000 steps per simulation.")
    doc.add_paragraph("In this environment, the 'Perceived Value of Society' (V_res) acts as a dynamic, shared resource representing general prosperity. During each step, both actors independently choose to act 'Hawkish' (aggressive/escalatory) or 'Dovish' (accommodating/de-escalatory). Hawkish actions extract a higher cost from the societal pool. Early termination of the simulation is triggered when systemic costs exhaust the total available value (PERCEIVED_VALUE = 10.0), indicating a structural societal failure (Revolution or Crackdown based on the distribution of pain).")
    
    doc.add_heading('2. Simulation Setup & Frameworks', level=1)
    doc.add_paragraph("The simulation sweeps across 8 distinct architectural frameworks to isolate the effects of differing strategic baselines and environmental sensitivities:")
    doc.add_paragraph("1. Simulation 1 (Stochastic Baseline): Completely random probability allocations at every step to map pure systemic noise.")
    doc.add_paragraph("2. Simulation 2 (Blind Control): Fixed ideological quadrants (DD, DH, HD, HH, RR) established at the start. These agents are perfectly oblivious to the decay of the societal resource pool.")
    doc.add_paragraph("3. Simulation 3 (Neutral Players): Random base hawkishness [0, 1] penalized by conscious awareness of structural decay (Aversion Penalties).")
    doc.add_paragraph("4. Simulation 4 (Neutral TFT): Players reactively copy the other's previous move. (If opponent played Dove, P_Hawk drops to [0, 0.5). If opponent played Hawk, P_Hawk jumps to (0.5, 1]).")
    doc.add_paragraph("5. Simulation 5 (Aggressive Ruler): Ruler has a fixed base hawkishness >0.5; Ruled remains neutral [0, 1].")
    doc.add_paragraph("6. Simulation 6 (Aggressive Ruler TFT): Ruler remains strictly Aggressive (>0.5); Ruled drops the neutral posture and adopts Tit-for-Tat copying.")
    doc.add_paragraph("7. Simulation 7 (Aggressive Ruled): Ruler is neutral; Ruled has a fixed base hawkishness >0.5.")
    doc.add_paragraph("8. Simulation 8 (Aggressive Ruled TFT): Ruler adopts TFT copying; Ruled remains strictly Aggressive (>0.5).")
    
    doc.add_heading('Strategic National Archetypes (Conscious Decay)', level=2)
    doc.add_paragraph("For simulations incorporating environmental decay awareness (Sims 3-8), we define five explicit aversion reactions:")
    doc.add_paragraph("1. Equals: Both Ruler and Ruled de-escalate collaboratively as resource pools drop (Linear 1:1 scaling).")
    doc.add_paragraph("2. Cowards: Symmetrically risk-averse, capitulating immediately into extreme dovish behavior under pressure (Quadratic scaling).")
    doc.add_paragraph("3. Fools: Symmetrically aggressive, resisting compromise and triggering stalemates (Sub-linear square-root scaling).")
    doc.add_paragraph("4. Tyrants: Hawkish Ruler (cushioned decay) and dovish Ruled (linear decay).")
    doc.add_paragraph("5. Brinksmen: Hawkish Ruled (cushioned decay) and dovish Ruler (linear decay).")
    
    doc.add_heading('3. Methodology & Data Processing', level=1)
    
    doc.add_heading('Principal Component Analysis (PCA)', level=2)
    doc.add_paragraph("Principal Component Analysis (PCA) is a robust dimensionality reduction algorithm used to visualize complex, multi-dimensional data architectures. In this analysis, we utilize PCA to compress six core behavioral metrics into a readable 2D geometric coordinate space:")
    
    # Isolate bullet points to apply explicit Left Alignment styles
    bullet_list = [
        "• Average Ruler Hawkishness",
        "• Average Ruled Hawkishness",
        "• Average Ruler Cost",
        "• Average Ruled Cost",
        "• Average Ruler Aversion Penalty",
        "• Average Ruled Aversion Penalty"
    ]
    for bullet_text in bullet_list:
        bullet_p = doc.add_paragraph(bullet_text)
        bullet_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    doc.add_paragraph("By explicitly excluding the actual 'Prosperity Score' (Cumulative V_res) from the PCA input matrix, we prevent scale-drowning and prevent the PCA from simply sorting by outcome. This approach ensures that the geometric clusters uniquely reveal the underlying behavioral manifolds, allowing us to observe how differing strategic postures naturally fall into high-prosperity or catastrophic outcome basins.")
    
    doc.add_heading('KDE Valley Clustering & Bandwidth Optimization', level=2)
    doc.add_paragraph("To objectively segment these geometric outcomes without relying on arbitrary binning or hardcoded thresholds, we perform a 1D Kernel Density Estimation (KDE) over the sorted Cumulative V_res scores. Finding the optimal curve bandwidth dynamically ensures that we capture genuine multi-modal outcome clusters. The algorithm sweeps through bandwidth multipliers at fine increments (0.01) to identify the model that preserves maximum structural granularity (valleys). To reject high-frequency noise and prevent overfitting, the system automatically excludes the top 5 sharpest peaks. The local minima ('valleys') in this optimized density distribution act as natural boundaries, defining discrete performance clusters.")
    
    doc.add_heading('Visual Marker Guide', level=2)
    doc.add_paragraph("Throughout the PCA scatter plots, distinct visual markers highlight systemic extremes within each space:")
    doc.add_paragraph("• Standard Survivals: Transparent circles with no borders.")
    doc.add_paragraph("• Failed Societies: Solid crosses, matching the cluster color, indicating Revolution or Crackdown.")
    doc.add_paragraph("• Elite Performers (Top 100): Highlighted with transparent diamonds featuring solid black borders, tracking the most prosperous trajectories.")
    doc.add_paragraph("• Worst Performers (Bottom 100): Highlighted with standard black crosses if the simulation failed, or transparent circles with a thin black border if they miraculously survived despite operating at the absolute bottom of the systemic prosperity curve.")

    doc.add_heading('4. State Failure Probability Analysis', level=1)
    doc.add_paragraph("This section details the exact probabilities of early state failure (and the corresponding triggers) across the strategic national groups. Each archetype is presented in its own table, with the blind Sim 2 Control included in every table as the comparative baseline.")
    
    archetypes = ['Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    df_control = df[df['Sim_Type'] == 2]
    
    def get_failure_stats(sub_df):
        total_n = len(sub_df)
        if total_n > 0:
            survived_count = len(sub_df[sub_df['Terminal_Status'] == 'Survived'])
            rev_count = len(sub_df[sub_df['Terminal_Status'] == 'Revolution'])
            crack_count = len(sub_df[sub_df['Terminal_Status'] == 'Crackdown'])
            prob_failure = (total_n - survived_count) / total_n
            prob_rev = rev_count / total_n
            prob_crack = crack_count / total_n
        else:
            prob_failure = prob_rev = prob_crack = 0.0
        return prob_failure, prob_rev, prob_crack

    ctrl_fail, ctrl_rev, ctrl_crack = get_failure_stats(df_control)
    
    for arch in archetypes:
        doc.add_heading(f"Archetype: {arch}", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text = 'Configuration', 'Total Failure'
        hdr_cells[2].text, hdr_cells[3].text = 'Revolution (Ruler Lost)', 'Crackdown (Ruled Lost)'
        
        # Insert Control
        row_cells = table.add_row().cells
        row_cells[0].text = 'Sim 2 Control'
        row_cells[1].text = f"{ctrl_fail:.2%}"
        row_cells[2].text = f"{ctrl_rev:.2%}"
        row_cells[3].text = f"{ctrl_crack:.2%}"
        
        for sim_type in range(3, 9):
            sname = {3:'Neutral', 4:'Neutral TFT', 5:'Agg Ruler', 6:'Agg Ruler TFT', 7:'Agg Ruled', 8:'Agg Ruled TFT'}[sim_type]
            sub_df = df[(df['Sim_Type'] == sim_type) & (df['Archetype'] == arch)]
            p_fail, p_rev, p_crack = get_failure_stats(sub_df)
            
            row_cells = table.add_row().cells
            row_cells[0].text = f"Sim {sim_type} {sname}"
            row_cells[1].text = f"{p_fail:.2%}"
            row_cells[2].text = f"{p_rev:.2%}"
            row_cells[3].text = f"{p_crack:.2%}"

    doc.add_heading('5. Analytical Visualizations', level=1)
    
    doc.add_heading('5.1 Unified PCA Projections', level=2)
    doc.add_paragraph("The Unified PCA plot projects all national archetypes together into a single global 2D manifold. This visualization helps us identify how differing strategic behavioral groups separate and cluster globally based on their core performance variables.")
    add_centered_image(doc, images[0], Inches(6.0))
    
    doc.add_heading('5.2 Case-Specific Local PCA Projections', level=2)
    doc.add_paragraph("These projections isolate the fixed starting quadrants of the Sim 2 Control group.")
    
    cases = ['DD', 'DH', 'HD', 'HH', 'RR']
    for i, case in enumerate(cases):
        doc.add_heading(f"Case {case} Localized Space", level=3)
        add_centered_image(doc, images[1 + i], Inches(5.0))
    
    doc.add_heading('5.3 Category-Specific Local PCA Projections', level=2)
    doc.add_paragraph("These category-specific projections isolate each individual framework. The PCA and scaling here are calculated locally strictly on the sliced data, revealing internal variance, survival thresholds, and fine-grained cluster boundaries specific to that archetype's environment.")
    doc.add_heading("Sim 2 Control Baseline Space", level=3)
    add_centered_image(doc, rendered_paths[('category', 'Sim_2_Control')], Inches(5.0))
    
    for arch in archetypes:
        doc.add_heading(f"Archetype: {arch}", level=3)
        doc.add_paragraph("Top Row: Neutral Players (Base vs TFT) | Middle Row: Aggressive Ruler (Base vs TFT) | Bottom Row: Aggressive Ruled (Base vs TFT)")
        
        img_table = doc.add_table(rows=3, cols=2)
        pairs = [
            (f"Sim_3_{arch}", f"Sim_4_{arch}"),
            (f"Sim_5_{arch}", f"Sim_6_{arch}"),
            (f"Sim_7_{arch}", f"Sim_8_{arch}")
        ]
        
        for r_idx, (left_key, right_key) in enumerate(pairs):
            cell_left = img_table.cell(r_idx, 0)
            cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_left = cell_left.paragraphs[0].add_run()
            run_left.add_picture(rendered_paths[('category', left_key)], width=Inches(3.0))
            
            cell_right = img_table.cell(r_idx, 1)
            cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_right = cell_right.paragraphs[0].add_run()
            run_right.add_picture(rendered_paths[('category', right_key)], width=Inches(3.0))
            
    doc.add_heading('5.4 Cluster Abundance Sweeps', level=2)
    doc.add_paragraph("These plots demonstrate the un-filtered KDE cluster population density across each simulation type.")
    for i in range(8):
        add_centered_image(doc, images[6 + i], Inches(5.5))
        
    doc.add_heading('5.5 Stacked Early State Failure Rates', level=2)
    doc.add_paragraph("These stacked bar charts isolate the exact modes of failure—Revolution or Crackdown. By excluding survivals, the heights of the bars represent the total proportional probability that a given configuration will fail. Sample sizes (N=...) are provided on the x-axis to confirm distribution consistency.")
    for img_path in images[-3:]:
        add_centered_image(doc, img_path, Inches(6.0))
        
    doc.add_heading('5.6 4-Panel Performance Metrics', level=2)
    doc.add_paragraph("This 4-panel breakdown separates the underlying variables—Ruler Hawkishness, Ruled Hawkishness, Ruler Cost, and Ruled Cost—for both the top performing and bottom performing states. This structure isolates the precise operational behaviors that drive the best (and worst) outcomes within each archetype.")
    doc.add_paragraph("Note on Negative Costs: The simulation calculates costs such that accommodating (Dovish) behavior actively reduces the cost burden. Consequently, performance metrics occasionally yield negative average costs, particularly among highly cooperative top performers. In these charts, negative cost indicates net added value (prosperity preservation) rather than debt.")
    add_centered_image(doc, perf_images[0], Inches(6.0)) # Top Performers
    add_centered_image(doc, perf_images[1], Inches(6.0)) # Bottom Performers
    
    doc.add_heading('5.7 Consolidated Prosperity Score Comparison', level=2)
    doc.add_paragraph("This final comparison chart aligns the absolute Prosperity Scores (Cumulative V_res) across all frameworks. Separated from the other metrics to prevent scale-drowning, it effectively compares the absolute wealth generation efficiency across strategies and archetypes.")
    add_centered_image(doc, perf_images[2], Inches(6.0)) # Prosperity
    
    doc.add_heading('6. Conclusion & Core Insights', level=1)
    doc.add_paragraph("Based on the comprehensive multi-framework analysis generated by this simulation run, several profound socio-structural dynamics emerge:")
    doc.add_paragraph("1. The Stabilizing Power of Tit-for-Tat (TFT): Across nearly all aggressive and unyielding archetypes (most notably the Fools and Tyrants), the introduction of a reactive, copy-based strategy dramatically curtails systemic collapse. By structurally guaranteeing retaliation against aggressive overreach, TFT functions as a powerful institutional governor, breaking runaway escalatory cycles before they drain the societal resource pool.")
    doc.add_paragraph("2. The Lethality of Strategic Blindness: The Sim 2 Control baseline operates with fixed ideological parameters and zero internal awareness of the wasting systemic value (V_res). This configuration consistently yields the highest rates of catastrophic collapse. This suggests that structural ignorance or political blindness to environmental/systemic decay is inherently more fatal to a society than explicit aggression.")
    doc.add_paragraph("3. The Preservation Efficiency of Symmetrical De-escalation: The 'Cowards' (highly risk-averse) and 'Equals' (cooperative) archetypes consistently demonstrate near-perfect survival rates and maximize the prosperity metric. In a multi-round environment characterized by mutual destruction (the Hawk-Dove dynamic), rapid, symmetrical capitulation in response to resource scarcity effectively insulates the collective wealth from structural friction.")

    report_path = os.path.join(target_dir, FILE_REPORT)
    doc.save(report_path)
    return report_path

def create_archive(df_res, df_hist, df_pca_feat, report_path, plots_dir):
    global global_actual_bw
    counter = 1
    while os.path.exists(f"{OUTPUT_PREFIX}_{counter}.zip"):
        counter += 1
    zip_filename = f"{OUTPUT_PREFIX}_{counter}.zip"
    

    print(f"\rPreparing optimized export data sheets... ")


    summary_xlsx = os.path.join(plots_dir, FILE_SUMMARY)
    history_xlsx = os.path.join(plots_dir, FILE_HISTORY)
    pca_xlsx = os.path.join(plots_dir, FILE_PCA_FEATURES)
    params_txt = os.path.join(plots_dir, FILE_PARAMS)
    
    print("Preparing Results Spread Sheet...")
    df_res.to_excel(summary_xlsx, index=False)

    print("Preparing PCA Spread Sheet...")
    df_pca_feat.to_excel(pca_xlsx, index=False)
    
    
    if not df_hist.empty:
        print("Preparing History Spread Sheet...")
        df_hist_str = df_hist.map(lambda x: str(x) if pd.notna(x) else "")
        df_hist_str.to_excel(history_xlsx)
    else:
        pd.DataFrame({"Status": ["History logging disabled."]}).to_excel(history_xlsx, index=False)

    print("Preparing 'params.txt'...")
    with open(params_txt, "w") as f:
        f.write(f"PERCEIVED_VALUE = {PERCEIVED_VALUE}\n")
        f.write(f"MAX_COST_PER_STEP = {MAX_COST_PER_STEP}\n")
        f.write(f"MAX_STEPS = {MAX_STEPS}\n")
        f.write(f"N_REPS_PER_CONFIG = {N_REPS}\n")
        f.write(f"KDE_BANDWIDTH_ADJUST = {global_actual_bw} (Mode: {KDE_BW_ADJUST})\n")
        f.write(f"KDE_TOP_N_PEAKS = {KDE_TOP_N_PEAKS}\n")
        f.write(f"KDE_EXCLUDE_TOP_N_PEAKS = {KDE_EXCLUDE_TOP_N_PEAKS}\n")
        f.write(f"KDE_OPTIMISATION_STEP = {KDE_OPTIMISATION_STEP}\n")
        f.write(f"PLOT_DOT_SIZE = {PLOT_DOT_SIZE}\n")

    print("Preparing Graph Folder...")
    files_to_pack = [(report_path, FILE_REPORT), (summary_xlsx, FILE_SUMMARY), (history_xlsx, FILE_HISTORY), (pca_xlsx, FILE_PCA_FEATURES), (params_txt, FILE_PARAMS)]
    for f in os.listdir(plots_dir):
        if f.endswith(".png"): files_to_pack.append((os.path.join(plots_dir, f), os.path.join("Graphs", f)))


    total_pack_files = len(files_to_pack)
    print(f"Packaging {total_pack_files} files directly into {zip_filename}...")
    start_pack_time = time.time()
    
    with zipfile.ZipFile(zip_filename, 'w') as zipf:
        for idx, (filepath, arcname) in enumerate(files_to_pack):
            zipf.write(filepath, arcname=arcname)
            completed = idx + 1
            elapsed = time.time() - start_pack_time
            speed = completed / elapsed if elapsed > 0 else 0
            bar_length = 30
            filled_length = int(round(bar_length * completed / total_pack_files))
            bar_chars = '#' * filled_length + '-' * (bar_length - filled_length)
            percentage = (completed / total_pack_files) * 100
            print(f"\rPackaging ZIP Archive: [{bar_chars}] {completed}/{total_pack_files} ({percentage:.1f}%) | Elapsed: {elapsed:.1f}s | Speed: {speed:.1f} files/s", end="", flush=True)
            
    print("\nCleaning up intermediate workspace files...")
    shutil.rmtree(plots_dir)
    print(f"Archive successfully generated: {zip_filename}")

# ==========================================
# 9. RUNNING THE BATCH SIMULATIONS
# ==========================================
def run_all_simulations():
    cases = ['DD', 'DH', 'HD', 'HH', 'RR']
    archetypes = ['Equals', 'Cowards', 'Fools', 'Brinksmen', 'Tyrants']
    jobs = []
    
    # Sim 1 (10x N_REPS)
    for _ in range(N_REPS * 10):
        jobs.append(('SIMULATE', (1, None, None)))
        
    # Sim 2 (N_REPS per case)
    for case in cases:
        for _ in range(N_REPS):
            jobs.append(('SIMULATE', (2, case, None)))
            
    # Sims 3 through 8 (Standardized identical simulation volume per strategic group)
    # Passed None for case here to strictly isolate 'RR' to Sim 2
    for sim_type in range(3, 9):
        for arch in archetypes:
            for _ in range(N_REPS * 5): 
                jobs.append(('SIMULATE', (sim_type, None, arch)))
                
    total_tasks = len(jobs)
    print(f"Total simulations to execute: {total_tasks}")
    task_results = run_parallel_jobs(jobs, desc="Executing Simulations")
    
    results = []
    histories = {}
    for idx, (res, hist) in enumerate(task_results):
        results.append(res)
        if RECORD_HISTORY:
            histories[idx] = hist
            
    df_res = pd.DataFrame(results)
    df_hist = pd.DataFrame() 
    return df_res, df_hist

# ==========================================
# 10. MAIN EXECUTION ROUTINE
# ==========================================
if __name__ == "__main__":
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    os.makedirs(TEMP_DIR)
    
    init_persistent_workers()
    try:
        df_res, df_hist = run_all_simulations()
        df_res, df_pca_feat, top_5, bottom_5_surv, term_clusters = process_data(df_res)
        images, rendered_paths = generate_graphs(df_res, top_5, bottom_5_surv, term_clusters, TEMP_DIR)
        perf_images = generate_performance_graphs(df_res, TEMP_DIR)
        report_path = generate_report(images, rendered_paths, perf_images, TEMP_DIR, df_res)
        create_archive(df_res, df_hist, df_pca_feat, report_path, TEMP_DIR)
    finally:
        shutdown_persistent_workers()
